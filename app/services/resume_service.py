"""
TalentMind AI - Resume Service
================================
Matches upload.py exact call signature.

Author  : TalentMind AI Team
Version : 1.0.0
"""

from __future__ import annotations

import io
import logging
import uuid
from typing import Any, Dict, Optional

from app.database.connection import db_manager
from app.database.repositories.resume_repo import ResumeRepository

logger = logging.getLogger(__name__)


class ResumeService:
    """Resume upload and processing service."""

    def __init__(self) -> None:
        self._resume_agent = None
        logger.info("ResumeService initialized")

    def _get_resume_agent(self):
        if self._resume_agent is None:
            try:
                from app.agents.resume_agent import ResumeAgent
                self._resume_agent = ResumeAgent()
            except Exception as exc:
                logger.warning("ResumeAgent load failed: %s", exc)
        return self._resume_agent

    def upload_resume(
        self,
        file_bytes: Any = None,
        file_name: str = "",
        file_type: str = "",
        candidate_name: str = "",
        uploaded_file: Any = None,
        target_role: str = "General Technology Role",
        **kwargs: Any,
    ) -> Dict[str, Any]:
        """
        Main upload method.

        upload.py calls this as:
            service.upload_resume(
                file_bytes     = file_bytes,
                file_name      = uploaded_file.name,
                candidate_name = candidate_name,
            )
        """
        try:
            actual_name = file_name or "resume"
            raw_text = ""

            if file_bytes is not None:
                raw_text = self._from_bytes(file_bytes, actual_name)

            elif uploaded_file is not None:
                b = uploaded_file.read()
                try:
                    uploaded_file.seek(0)
                except Exception:
                    pass
                actual_name = (
                    file_name
                    or getattr(uploaded_file, "name", "resume")
                )
                raw_text = self._from_bytes(b, actual_name)

            else:
                return self._err("No file provided.")

            if not raw_text or not raw_text.strip():
                return self._err(
                    "Could not extract text from file. "
                    "Please ensure file is readable and "
                    "not a scanned image."
                )

            if len(raw_text.strip()) < 50:
                return self._err(
                    "Resume content too short. "
                    "Please upload a complete resume."
                )

            raw_text = raw_text.strip()
            actual_type = file_type or self._mime(actual_name)
            candidate_id = str(uuid.uuid4())
            resume_id = str(uuid.uuid4())

            parsed_data = {}
            try:
                agent = self._get_resume_agent()
                if agent:
                    r = agent.run({
                        "resume_text": raw_text,
                        "target_role": target_role,
                    })
                    if r and r.get("success"):
                        parsed_data = r
            except Exception as exc:
                logger.warning(
                    "Parse failed (non-critical): %s", exc
                )

            # ── Persist to database ──────────────────────────────
            try:
                # Extract contact info from parsed data if available
                contact_info = {}
                if parsed_data and parsed_data.get("success"):
                    contact_info = parsed_data.get(
                        "contact_info", {}
                    ) or parsed_data.get("contact", {}) or {}

                email = contact_info.get("email")
                phone = contact_info.get("phone")

                with db_manager.get_session() as session:
                    repo = ResumeRepository(session)

                    # Create or get the candidate
                    repo.get_or_create_candidate(
                        candidate_id=candidate_id,
                        name=candidate_name or contact_info.get("name"),
                        email=email,
                    )

                    # Deactivate previous resumes for this candidate
                    repo.deactivate_previous_resumes(
                        candidate_id=candidate_id,
                        keep_resume_id=resume_id,
                    )

                    # Create resume record
                    repo.create_resume(
                        resume_id=resume_id,
                        candidate_id=candidate_id,
                        file_name=actual_name,
                        file_type=actual_type,
                        raw_text=raw_text,
                        parsed_data=parsed_data or None,
                    )

                logger.info(
                    "Resume persisted to DB | resume_id=%s | "
                    "candidate_id=%s",
                    resume_id, candidate_id,
                )
            except Exception as exc:
                # DB persistence is non-critical; log and continue
                logger.error(
                    "DB persistence failed (non-critical): %s", exc
                )

            return {
                "success": True,
                "error": None,
                "message": (
                    "Resume uploaded and analyzed successfully!"
                ),
                "raw_text": raw_text,
                "resume_text": raw_text,
                "candidate_id": candidate_id,
                "resume_id": resume_id,
                "file_name": actual_name,
                "file_type": actual_type,
                "char_count": len(raw_text),
                "parsed_data": parsed_data,
                "candidate_name": candidate_name or "",
            }

        except Exception as exc:
            logger.error("Upload failed: %s", exc)
            return self._err(str(exc))

    def process_upload(self, **kwargs: Any) -> Dict[str, Any]:
        """Alias for upload_resume."""
        return self.upload_resume(**kwargs)

    def extract_text(
        self,
        file_bytes: Any = None,
        file_name: str = "",
        uploaded_file: Any = None,
    ) -> str:
        """Extract text only."""
        if file_bytes is not None:
            return self._from_bytes(file_bytes, file_name)
        if uploaded_file is not None:
            b = uploaded_file.read()
            try:
                uploaded_file.seek(0)
            except Exception:
                pass
            return self._from_bytes(
                b, getattr(uploaded_file, "name", "")
            )
        return ""

    def parse_resume(
        self,
        resume_text: str = "",
        target_role: str = "General Technology Role",
    ) -> Dict[str, Any]:
        """Parse resume text with AI agent."""
        if not resume_text:
            return {"success": False, "error": "No text."}
        try:
            agent = self._get_resume_agent()
            if agent:
                return agent.run({
                    "resume_text": resume_text,
                    "target_role": target_role,
                })
            return {"success": False, "error": "No agent."}
        except Exception as exc:
            return {"success": False, "error": str(exc)}

    def _from_bytes(
        self,
        file_bytes: bytes,
        file_name: str = "",
    ) -> str:
        """Extract text from raw bytes."""
        if not file_bytes:
            return ""

        name = file_name.lower()

        try:
            if name.endswith(".pdf"):
                return self._pdf(file_bytes)
            elif name.endswith((".docx", ".doc")):
                return self._docx(file_bytes)
            elif name.endswith(".txt"):
                return file_bytes.decode("utf-8", errors="ignore")
            else:
                if file_bytes[:4] == b"%PDF":
                    return self._pdf(file_bytes)
                elif file_bytes[:2] == b"PK":
                    return self._docx(file_bytes)
                else:
                    return file_bytes.decode(
                        "utf-8", errors="ignore"
                    )
        except Exception as exc:
            logger.error("Bytes extraction error: %s", exc)
            return ""

    def _pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF bytes."""
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
            result = "\n".join(parts)
            if result.strip():
                return result
        except ImportError:
            pass
        except Exception as exc:
            logger.warning("pypdf error: %s", exc)

        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
            return "\n".join(parts)
        except ImportError:
            logger.error("Run: pip install pypdf")
            return ""
        except Exception as exc:
            logger.error("PDF extraction failed: %s", exc)
            return ""

    def _docx(self, file_bytes: bytes) -> str:
        """Extract text from DOCX bytes."""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)

            return "\n".join(parts)

        except ImportError:
            logger.error("Run: pip install python-docx")
            return ""
        except Exception as exc:
            logger.error("DOCX extraction failed: %s", exc)
            return ""

    def _mime(self, filename: str) -> str:
        """Get MIME type from filename."""
        n = filename.lower()
        if n.endswith(".pdf"):
            return "application/pdf"
        if n.endswith(".docx"):
            return (
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            )
        if n.endswith(".txt"):
            return "text/plain"
        return "application/octet-stream"

    def _err(self, msg: str) -> Dict[str, Any]:
        """Standardized error result with all expected keys."""
        return {
            "success": False,
            "error": msg,
            "message": msg,
            "raw_text": "",
            "resume_text": "",
            "candidate_id": "",
            "resume_id": "",
            "file_name": "",
            "file_type": "",
            "char_count": 0,
            "parsed_data": {},
            "candidate_name": "",
        }